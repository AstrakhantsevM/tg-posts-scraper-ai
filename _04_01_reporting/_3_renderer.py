"""
_04_01_report/_3_docx_renderer.py

Рендеринг ReportDocument в профессиональный Word-документ (.docx)
в стиле «Мониторинг СМИ».

Исправленная UX/UI-версия:
- шапка во всю ширину страницы;
- крупный читаемый заголовок;
- горизонтальные счётчики-виджеты справа;
- карточки регионов во всю ширину;
- каждая мера поддержки выводится отдельной мини-карточкой;
- title/type/name корректно выводятся как суть меры;
- строковые null/None/не указано не попадают в документ;
- настоящий футер внизу каждой страницы;
- номер страницы в футере;
- подпись «Отчёт сформирован автоматически» на каждой странице.
"""

from __future__ import annotations

import re
from datetime import datetime, timezone

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor, Twips

from _04_01_reporting._0_models import RegionReport, ReportDocument


# ---------------------------------------------------------------------------
# Цветовая палитра
# ---------------------------------------------------------------------------

COLOR_NAVY = RGBColor(0x17, 0x34, 0x5E)
COLOR_BLUE = RGBColor(0x2F, 0x6F, 0xB2)
COLOR_BLUE_SOFT = "EAF2FF"

COLOR_GREEN = RGBColor(0x5E, 0x9F, 0x3F)
COLOR_GREEN_BG = "EAF6E6"

COLOR_RED = RGBColor(0xC0, 0x2A, 0x2A)
COLOR_RED_BG = "FBEAEA"

COLOR_PURPLE = RGBColor(0x6D, 0x4A, 0xA8)
COLOR_PURPLE_BG = "F1ECFA"

COLOR_GREY = RGBColor(0x68, 0x68, 0x68)
COLOR_GREY_DARK = RGBColor(0x45, 0x45, 0x45)
COLOR_GREY_BG = "F4F6F8"

COLOR_LINE = "D7DEE8"
COLOR_CARD_BG = "FAFBFD"
COLOR_WHITE = "FFFFFF"


# ---------------------------------------------------------------------------
# Размеры шрифтов
# ---------------------------------------------------------------------------

F_NORMAL = 10.0
F_SMALL = 8.5
F_META = 8.8
F_HEADER = 20.0
F_REGION = 12.0
F_SECTION = 11.0
F_STAT = 15.0


# ---------------------------------------------------------------------------
# XML-хелперы
# ---------------------------------------------------------------------------

def _set_cell_bg(cell, fill_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex.upper())
    tc_pr.append(shd)


def _set_cell_borders(cell, *, color: str = "CCCCCC", size: int = 4) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()

    old_borders = tc_pr.find(qn("w:tcBorders"))
    if old_borders is not None:
        tc_pr.remove(old_borders)

    tc_borders = OxmlElement("w:tcBorders")

    for side in ("top", "bottom", "left", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color.upper())
        tc_borders.append(el)

    tc_pr.append(tc_borders)


def _cell_margins(cell, top=60, bottom=60, left=100, right=100) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()

    old_mar = tc_pr.find(qn("w:tcMar"))
    if old_mar is not None:
        tc_pr.remove(old_mar)

    tc_mar = OxmlElement("w:tcMar")

    for side, val in [
        ("top", top),
        ("bottom", bottom),
        ("left", left),
        ("right", right),
    ]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        tc_mar.append(el)

    tc_pr.append(tc_mar)


def _p_border_bottom(p, color_hex: str = "C0C0C0", size: int = 4) -> None:
    p_pr = p._p.get_or_add_pPr()

    old_bdr = p_pr.find(qn("w:pBdr"))
    if old_bdr is not None:
        p_pr.remove(old_bdr)

    p_bdr = OxmlElement("w:pBdr")

    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color_hex.upper())

    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _set_line_spacing(p, line: int = 240) -> None:
    """
    Одинарный межстрочный интервал: 240 = 1.0×.
    """
    p_pr = p._p.get_or_add_pPr()

    old_spacing = p_pr.find(qn("w:spacing"))
    if old_spacing is not None:
        p_pr.remove(old_spacing)

    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:line"), str(line))
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)


def _clear_cell(cell) -> None:
    for p in list(cell.paragraphs):
        p._p.getparent().remove(p._p)


def _set_table_fixed_layout(tbl) -> None:
    tbl.autofit = False

    tbl_pr = tbl._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))

    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)

    layout.set(qn("w:type"), "fixed")


def _set_tbl_width(tbl, width_dxa: int) -> None:
    tbl_pr = tbl._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))

    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)

    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")


def _remove_table_borders(tbl) -> None:
    tbl_pr = tbl._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))

    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)

    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(qn(f"w:{side}"))

        if el is None:
            el = OxmlElement(f"w:{side}")
            borders.append(el)

        el.set(qn("w:val"), "nil")


def _cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))

    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)

    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def _add_page_number(paragraph) -> None:
    run = paragraph.add_run()

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")

    text = OxmlElement("w:t")
    text.text = "1"

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(text)
    run._r.append(fld_end)

    run.font.size = Pt(F_SMALL)
    run.font.color.rgb = COLOR_GREY


def _override_list_bullet_style(doc: Document) -> None:
    """
    Убираем лишние интервалы в List Bullet.
    """
    try:
        style = doc.styles["List Bullet"]
        style.font.name = "Arial"
        style.font.size = Pt(F_NORMAL)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.left_indent = Twips(360)
        style.paragraph_format.first_line_indent = Twips(-180)
    except Exception:
        pass


# ---------------------------------------------------------------------------
# Рендерер
# ---------------------------------------------------------------------------

class DocxReportRenderer:
    """
    Превращает ReportDocument в компактный Word-документ для социологов.
    """

    PAGE_WIDTH = Mm(210)
    PAGE_HEIGHT = Mm(297)
    MARGIN = Mm(14)

    CONTENT_W = PAGE_WIDTH - 2 * MARGIN

    # 1 мм ≈ 56.7 twips.
    CONTENT_W_DXA = int((210 - 14 * 2) * 56.7)

    def render(self, document: ReportDocument) -> Document:
        doc = Document()

        self._setup_page(doc)
        self._setup_styles(doc)
        _override_list_bullet_style(doc)

        try:
            gen_dt = datetime.fromisoformat(document.generated_at)
        except ValueError:
            gen_dt = datetime.now(tz=timezone.utc)

        date_str = gen_dt.strftime("%d.%m.%Y")
        time_str = gen_dt.strftime("%H:%M")

        found_regions = [r for r in document.regions if r.found]
        empty_regions = [r for r in document.regions if not r.found and not r.errors]
        error_regions = [r for r in document.regions if r.errors]

        self._render_page_footer(doc, date_str, time_str)

        self._render_header(
            doc=doc,
            document=document,
            date_str=date_str,
            found_regions=found_regions,
            empty_regions=empty_regions,
            error_regions=error_regions,
        )

        if found_regions:
            self._render_section_title(
                doc,
                "ВЫЯВЛЕННЫЕ МЕРЫ СОЦИАЛЬНОЙ ПОДДЕРЖКИ",
            )

            for idx, region in enumerate(found_regions, start=1):
                self._render_region_found(doc, idx, region)
        else:
            self._render_no_mentions_message(doc)

        self._render_empty_footer(doc, empty_regions, error_regions)

        return doc

    # ------------------------------------------------------------------
    # Базовая настройка документа
    # ------------------------------------------------------------------

    def _setup_page(self, doc: Document) -> None:
        for section in doc.sections:
            section.page_width = self.PAGE_WIDTH
            section.page_height = self.PAGE_HEIGHT

            section.top_margin = Mm(13)
            section.bottom_margin = Mm(18)
            section.left_margin = self.MARGIN
            section.right_margin = self.MARGIN

            section.footer_distance = Mm(7)

    def _setup_styles(self, doc: Document) -> None:
        normal = doc.styles["Normal"]
        normal.font.name = "Arial"
        normal.font.size = Pt(F_NORMAL)
        normal.paragraph_format.space_before = Pt(0)
        normal.paragraph_format.space_after = Pt(0)

    # ------------------------------------------------------------------
    # Футер
    # ------------------------------------------------------------------

    def _render_page_footer(self, doc: Document, date_str: str, time_str: str) -> None:
        for section in doc.sections:
            footer = section.footer
            footer.is_linked_to_previous = False

            first_p = footer.paragraphs[0]
            first_p.text = ""
            first_p.paragraph_format.space_before = Pt(0)
            first_p.paragraph_format.space_after = Pt(0)

            tbl = footer.add_table(rows=1, cols=2, width=self.CONTENT_W)
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

            _set_table_fixed_layout(tbl)
            _remove_table_borders(tbl)
            _set_tbl_width(tbl, self.CONTENT_W_DXA)

            left = tbl.rows[0].cells[0]
            right = tbl.rows[0].cells[1]

            _cell_width(left, int(self.CONTENT_W_DXA * 0.74))
            _cell_width(right, int(self.CONTENT_W_DXA * 0.26))

            _cell_margins(left, top=30, bottom=0, left=0, right=0)
            _cell_margins(right, top=30, bottom=0, left=0, right=0)

            lp = left.paragraphs[0]
            lp.paragraph_format.space_before = Pt(0)
            lp.paragraph_format.space_after = Pt(0)
            _p_border_bottom(lp, color_hex="E1E6EF", size=2)

            run = lp.add_run(
                f"Отчёт сформирован автоматически · {date_str} {time_str}"
            )
            run.italic = True
            run.font.size = Pt(F_SMALL)
            run.font.color.rgb = COLOR_GREY

            rp = right.paragraphs[0]
            rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            rp.paragraph_format.space_before = Pt(0)
            rp.paragraph_format.space_after = Pt(0)
            _p_border_bottom(rp, color_hex="E1E6EF", size=2)

            r = rp.add_run("стр. ")
            r.font.size = Pt(F_SMALL)
            r.font.color.rgb = COLOR_GREY

            _add_page_number(rp)

    # ------------------------------------------------------------------
    # Шапка
    # ------------------------------------------------------------------

    def _render_header(
        self,
        doc: Document,
        document: ReportDocument,
        date_str: str,
        found_regions: list[RegionReport],
        empty_regions: list[RegionReport],
        error_regions: list[RegionReport],
    ) -> None:
        tbl = doc.add_table(rows=1, cols=2)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.style = "Table Grid"

        _set_table_fixed_layout(tbl)
        _set_tbl_width(tbl, self.CONTENT_W_DXA)
        _remove_table_borders(tbl)

        left_w = int(self.CONTENT_W_DXA * 0.72)
        right_w = self.CONTENT_W_DXA - left_w

        cell_l = tbl.rows[0].cells[0]
        cell_r = tbl.rows[0].cells[1]

        _cell_width(cell_l, left_w)
        _cell_width(cell_r, right_w)

        _set_cell_bg(cell_l, COLOR_BLUE_SOFT)
        _set_cell_bg(cell_r, COLOR_BLUE_SOFT)

        _cell_margins(cell_l, top=180, bottom=160, left=220, right=160)
        _cell_margins(cell_r, top=130, bottom=120, left=120, right=180)

        cell_l.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell_r.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        p_title = cell_l.paragraphs[0]
        p_title.paragraph_format.space_before = Pt(0)
        p_title.paragraph_format.space_after = Pt(4)
        _set_line_spacing(p_title)

        run = p_title.add_run("МОНИТОРИНГ СМИ")
        run.bold = True
        run.font.size = Pt(F_HEADER)
        run.font.color.rgb = COLOR_NAVY

        p_sub = cell_l.add_paragraph()
        p_sub.paragraph_format.space_before = Pt(0)
        p_sub.paragraph_format.space_after = Pt(0)
        _set_line_spacing(p_sub)

        rs = p_sub.add_run("Итоговый отчёт")
        rs.bold = True
        rs.font.size = Pt(F_META)
        rs.font.color.rgb = COLOR_GREY_DARK

        rm = p_sub.add_run(
            f"  ·  {date_str}  ·  регионов: {document.regions_total}"
        )
        rm.font.size = Pt(F_META)
        rm.font.color.rgb = COLOR_GREY

        _clear_cell(cell_r)

        stats = [
            (str(len(found_regions)), "найдено", COLOR_GREEN, COLOR_GREEN_BG),
            (str(len(error_regions)), "ошибок", COLOR_RED, COLOR_RED_BG),
            (str(len(empty_regions)), "пусто", COLOR_GREY, COLOR_GREY_BG),
        ]

        inner = cell_r.add_table(rows=1, cols=3)
        inner.alignment = WD_TABLE_ALIGNMENT.RIGHT

        _set_table_fixed_layout(inner)
        _remove_table_borders(inner)

        usable_right_w = right_w - 120 - 180
        stat_w = int(usable_right_w / 3)
        _set_tbl_width(inner, usable_right_w)

        for i, (num, label, color, bg) in enumerate(stats):
            c = inner.rows[0].cells[i]

            _cell_width(c, stat_w)
            _set_cell_bg(c, bg)
            _set_cell_borders(c, color="FFFFFF", size=4)
            _cell_margins(c, top=80, bottom=70, left=50, right=50)

            p_num = c.paragraphs[0]
            p_num.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_num.paragraph_format.space_before = Pt(0)
            p_num.paragraph_format.space_after = Pt(0)
            _set_line_spacing(p_num)

            rn = p_num.add_run(num)
            rn.bold = True
            rn.font.size = Pt(F_STAT)
            rn.font.color.rgb = color

            p_label = c.add_paragraph()
            p_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_label.paragraph_format.space_before = Pt(0)
            p_label.paragraph_format.space_after = Pt(0)
            _set_line_spacing(p_label)

            rl = p_label.add_run(label)
            rl.font.size = Pt(F_SMALL)
            rl.font.color.rgb = COLOR_GREY_DARK

        sp = doc.add_paragraph()
        sp.paragraph_format.space_before = Pt(0)
        sp.paragraph_format.space_after = Pt(10)

    # ------------------------------------------------------------------
    # Заголовок раздела
    # ------------------------------------------------------------------

    def _render_section_title(self, doc: Document, text: str) -> None:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(6)
        _set_line_spacing(p)
        _p_border_bottom(p, color_hex="8FB3DF", size=8)

        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(F_SECTION)
        run.font.color.rgb = COLOR_NAVY

    # ------------------------------------------------------------------
    # Сообщение, если ничего не найдено
    # ------------------------------------------------------------------

    def _render_no_mentions_message(self, doc: Document) -> None:
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.style = "Table Grid"

        _set_table_fixed_layout(tbl)
        _set_tbl_width(tbl, self.CONTENT_W_DXA)
        _remove_table_borders(tbl)

        cell = tbl.rows[0].cells[0]

        _cell_width(cell, self.CONTENT_W_DXA)
        _set_cell_bg(cell, COLOR_GREY_BG)
        _set_cell_borders(cell, color="E0E4EA", size=5)
        _cell_margins(cell, top=110, bottom=105, left=150, right=150)

        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        _set_line_spacing(p)

        run = p.add_run("Ни в одном регионе упоминаний не обнаружено.")
        run.italic = True
        run.font.size = Pt(F_NORMAL)
        run.font.color.rgb = COLOR_GREY_DARK

    # ------------------------------------------------------------------
    # Регион с упоминаниями
    # ------------------------------------------------------------------

    def _render_region_found(
        self,
        doc: Document,
        idx: int,
        region: RegionReport,
    ) -> None:
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.style = "Table Grid"

        _set_table_fixed_layout(tbl)
        _set_tbl_width(tbl, self.CONTENT_W_DXA)
        _remove_table_borders(tbl)

        cell = tbl.rows[0].cells[0]

        _cell_width(cell, self.CONTENT_W_DXA)
        _set_cell_bg(cell, COLOR_WHITE)
        _set_cell_borders(cell, color=COLOR_LINE, size=5)
        _cell_margins(cell, top=120, bottom=115, left=150, right=150)

        # вместо p_name / p_meta
        meta_parts = []

        if region.channel:
            meta_parts.append(str(region.channel))

        if region.posts_total:
            meta_parts.append(f"{region.posts_total} постов")

        if region.data_date:
            meta_parts.append(region.data_date)

            # Выводим заголовок и мету в первый абзац карточки
            p_head = cell.paragraphs[0]
            p_head.paragraph_format.space_before = Pt(0)
            p_head.paragraph_format.space_after = Pt(12)  # Небольшой отступ перед карточками мер
            _set_line_spacing(p_head)

            # 1. Номер
            rn = p_head.add_run(f"{idx:02d}. ")
            rn.bold = True
            rn.font.size = Pt(F_REGION)
            rn.font.color.rgb = COLOR_BLUE

            # 2. Регион
            rr = p_head.add_run(region.region)
            rr.bold = True
            rr.font.size = Pt(F_REGION)
            rr.font.color.rgb = COLOR_NAVY

            # 3. Метаинформация сразу за названием региона
            if meta_parts:
                # Разделитель с отступами
                r_sep = p_head.add_run("   ·   ")
                r_sep.font.size = Pt(F_META)
                r_sep.font.color.rgb = COLOR_GREY

                rm = p_head.add_run(" · ".join(meta_parts))
                rm.italic = True
                rm.font.size = Pt(F_META)
                rm.font.color.rgb = COLOR_GREY

        spacer = cell.add_paragraph()
        spacer.paragraph_format.space_after = Pt(0)
        # Задаем высоту самого шрифта в пустой строке
        spacer.add_run().font.size = Pt(8)

        self._render_measures_into_cell(cell, region)

        sp = doc.add_paragraph()
        sp.paragraph_format.space_before = Pt(0)
        sp.paragraph_format.space_after = Pt(8)

    def _render_measures_into_cell(self, cell, region: RegionReport) -> None:
        if region.measures:
            for measure in region.measures:
                if isinstance(measure, dict):
                    self._render_measure_dict_into_cell(cell, measure)
                else:
                    text = self._clean_measure_value(measure)
                    if text:
                        self._render_bullet_into_cell(cell, text)

            return

        if region.summary and region.summary != "Упоминаний не найдено.":
            for line in region.summary.splitlines():
                line = re.sub(r"^[#*\-•]+\s*", "", line.strip())

                if line:
                    self._render_bullet_into_cell(cell, line)

    def _render_measure_dict_into_cell(self, cell, m: dict) -> None:
        """
        Рендерит одну меру как отдельную мини-карточку.
        """
        title = self._clean_measure_value(
            m.get("title")
            or m.get("name")
            or m.get("type")
            or m.get("description")
        )

        form = self._clean_measure_value(m.get("form"))
        amount = self._clean_measure_value(m.get("amount"))
        conditions = self._clean_measure_value(m.get("conditions"))
        notes = self._clean_measure_value(m.get("notes"))
        details = self._clean_measure_value(m.get("details"))

        if not any([title, form, amount, conditions, notes, details]):
            return

        if not title:
            title = "Мера поддержки без уточнённого названия"

        card = cell.add_table(rows=1, cols=1)
        card.alignment = WD_TABLE_ALIGNMENT.CENTER
        card.style = "Table Grid"

        _set_table_fixed_layout(card)
        _remove_table_borders(card)

        card_w = int(self.CONTENT_W_DXA * 0.94)
        _set_tbl_width(card, card_w)

        card_cell = card.rows[0].cells[0]
        _cell_width(card_cell, card_w)
        _set_cell_bg(card_cell, COLOR_CARD_BG)
        _set_cell_borders(card_cell, color="E1E6EF", size=5)
        _cell_margins(card_cell, top=85, bottom=80, left=120, right=120)

        p_title = card_cell.paragraphs[0]
        p_title.paragraph_format.space_before = Pt(0)
        p_title.paragraph_format.space_after = Pt(3)
        _set_line_spacing(p_title)

        bullet = p_title.add_run("● ")
        bullet.bold = True
        bullet.font.size = Pt(F_NORMAL)
        bullet.font.color.rgb = COLOR_BLUE

        run_title = p_title.add_run(title)
        run_title.bold = True
        run_title.font.size = Pt(F_NORMAL + 0.5)
        run_title.font.color.rgb = COLOR_NAVY

        badges = []

        if form:
            badges.append(("Форма", form, COLOR_PURPLE, COLOR_PURPLE_BG))

        if amount:
            badges.append(("Размер", amount, COLOR_GREEN, COLOR_GREEN_BG))

        if badges:
            badges_tbl = card_cell.add_table(rows=1, cols=len(badges))
            badges_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

            _set_table_fixed_layout(badges_tbl)
            _remove_table_borders(badges_tbl)

            badge_w = int(card_w / max(len(badges), 1)) - 80
            _set_tbl_width(badges_tbl, badge_w * len(badges))

            for i, (label, value, color, bg) in enumerate(badges):
                bcell = badges_tbl.rows[0].cells[i]

                _cell_width(bcell, badge_w)
                _set_cell_bg(bcell, bg)
                _set_cell_borders(bcell, color="FFFFFF", size=4)
                _cell_margins(bcell, top=45, bottom=40, left=70, right=70)

                p = bcell.paragraphs[0]
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                _set_line_spacing(p)

                rl = p.add_run(f"{label}: ")
                rl.bold = True
                rl.font.size = Pt(F_SMALL)
                rl.font.color.rgb = color

                rv = p.add_run(value)
                rv.font.size = Pt(F_SMALL)
                rv.font.color.rgb = COLOR_GREY_DARK

        detail_rows = []

        if conditions:
            detail_rows.append(("Условия", conditions))

        if notes:
            detail_rows.append(("Нюансы", notes))

        if details:
            detail_rows.append(("Детали", details))

        for label, value in detail_rows:
            p = card_cell.add_paragraph()
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.left_indent = Twips(0)
            _set_line_spacing(p)

            rl = p.add_run(f"{label}: ")
            rl.bold = True
            rl.font.size = Pt(F_NORMAL)
            rl.font.color.rgb = COLOR_GREY_DARK

            rv = p.add_run(value)
            rv.font.size = Pt(F_NORMAL)
            rv.font.color.rgb = COLOR_GREY_DARK

        #spacer = cell.add_paragraph()
        #spacer.paragraph_format.space_before = Pt(0)
        #spacer.paragraph_format.space_after = Pt(5)

    def _render_bullet_into_cell(self, cell, text: str) -> None:
        p = cell.add_paragraph()
        p.paragraph_format.left_indent = Twips(260)
        p.paragraph_format.first_line_indent = Twips(-160)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        _set_line_spacing(p)

        bullet = p.add_run("• ")
        bullet.bold = True
        bullet.font.size = Pt(F_NORMAL)
        bullet.font.color.rgb = COLOR_BLUE

        run = p.add_run(text)
        run.font.size = Pt(F_NORMAL)
        run.font.color.rgb = COLOR_GREY_DARK

    # ------------------------------------------------------------------
    # Не найдено / ошибки
    # ------------------------------------------------------------------

    def _render_empty_footer(
        self,
        doc: Document,
        empty_regions: list[RegionReport],
        error_regions: list[RegionReport],
    ) -> None:
        if not empty_regions and not error_regions:
            return

        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.style = "Table Grid"

        _set_table_fixed_layout(tbl)
        _set_tbl_width(tbl, self.CONTENT_W_DXA)
        _remove_table_borders(tbl)

        cell = tbl.rows[0].cells[0]

        _cell_width(cell, self.CONTENT_W_DXA)
        _set_cell_bg(cell, COLOR_GREY_BG)
        _set_cell_borders(cell, color="E0E4EA", size=5)
        _cell_margins(cell, top=100, bottom=95, left=150, right=150)

        p_title = cell.paragraphs[0]
        p_title.paragraph_format.space_before = Pt(0)
        p_title.paragraph_format.space_after = Pt(4)
        _set_line_spacing(p_title)

        run = p_title.add_run("УПОМИНАНИЙ НЕ НАЙДЕНО")
        run.bold = True
        run.font.size = Pt(F_NORMAL)
        run.font.color.rgb = COLOR_GREY_DARK

        if empty_regions:
            p = cell.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            _set_line_spacing(p)

            rv = p.add_run(", ".join(r.region for r in empty_regions))
            rv.font.size = Pt(F_META)
            rv.font.color.rgb = COLOR_GREY

        if error_regions:
            p_e = cell.add_paragraph()
            p_e.paragraph_format.space_before = Pt(5)
            p_e.paragraph_format.space_after = Pt(0)
            _set_line_spacing(p_e)

            rl = p_e.add_run("Ошибки: ")
            rl.bold = True
            rl.font.size = Pt(F_META)
            rl.font.color.rgb = COLOR_RED

            rv = p_e.add_run(", ".join(r.region for r in error_regions))
            rv.font.size = Pt(F_META)
            rv.font.color.rgb = COLOR_RED

    # ------------------------------------------------------------------
    # Очистка значений
    # ------------------------------------------------------------------

    def _clean_measure_value(self, value) -> str | None:
        if value is None:
            return None

        text = str(value).strip()

        if not text:
            return None

        text = re.sub(r"\s+", " ", text)

        if text.lower() in {
            "null",
            "none",
            "nil",
            "нет",
            "не указано",
            "не указаны",
            "не указан",
            "n/a",
            "na",
            "—",
            "-",
        }:
            return None

        return text